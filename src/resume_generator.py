from datetime import datetime
from pathlib import Path
import re

from src.config import OUTPUTS_DIR


SECTION_ALIASES = {
    "education": ["教育", "教育背景", "学习经历"],
    "internship": ["实习", "实习经历", "工作经历", "实践经历"],
    "project": ["项目", "项目经历", "作品", "作品集"],
    "skills": ["技能", "专业技能", "技能证书", "工具"],
    "other": ["校园", "其他", "获奖", "社团", "志愿"],
}


def _split_lines(text: str) -> list[str]:
    return [line.strip() for line in (text or "").splitlines() if line.strip()]


def _find_section_lines(resume_text: str, aliases: list[str], max_lines: int = 10) -> list[str]:
    """从原简历中按标题附近提取内容；找不到时返回空列表，避免编造。"""
    lines = _split_lines(resume_text)
    for index, line in enumerate(lines):
        if any(alias in line for alias in aliases):
            collected = []
            for candidate in lines[index + 1 : index + 1 + max_lines]:
                if any(candidate.startswith(alias) or candidate == alias for group in SECTION_ALIASES.values() for alias in group):
                    break
                collected.append(candidate)
            return collected
    return []


def _pick_relevant_lines(resume_text: str, keywords: list[str], limit: int = 8) -> list[str]:
    lines = _split_lines(resume_text)
    picked = []
    for line in lines:
        if any(keyword and keyword.lower() in line.lower() for keyword in keywords):
            picked.append(line)
        if len(picked) >= limit:
            break
    return picked


def _to_bullet(line: str) -> str:
    line = re.sub(r"^[\-•*]\s*", "", line).strip()
    if not line:
        return ""
    return f"- {line}"


def _format_section(title: str, lines: list[str], fallback: str = "原简历中未明确拆分该模块，请保留原文核对后再补充。") -> list[str]:
    output = [f"## {title}"]
    if lines:
        output.extend([_to_bullet(line) for line in lines if _to_bullet(line)])
    else:
        output.append(f"- {fallback}")
    return output


def generate_resume_draft(resume_text: str, jd_text: str, analysis: dict | None = None) -> str:
    """生成保守版定制简历：只重组、强化原简历已有内容，不虚构经历。"""
    analysis = analysis or {}
    keywords = []
    keywords.extend(analysis.get("matched_keywords", []))
    keywords.extend(analysis.get("jd_high_frequency_requirements", []))
    keywords.extend([item.get("jd_requirement", "") for item in analysis.get("matched_evidence", []) if isinstance(item, dict)])
    keywords = [str(keyword).strip() for keyword in keywords if str(keyword).strip()]

    relevant_lines = _pick_relevant_lines(resume_text, keywords, limit=8)
    education = _find_section_lines(resume_text, SECTION_ALIASES["education"], max_lines=6)
    internships = _find_section_lines(resume_text, SECTION_ALIASES["internship"], max_lines=12)
    projects = _find_section_lines(resume_text, SECTION_ALIASES["project"], max_lines=12)
    skills = _find_section_lines(resume_text, SECTION_ALIASES["skills"], max_lines=8)
    other = _find_section_lines(resume_text, SECTION_ALIASES["other"], max_lines=8)

    strengths = analysis.get("strengths", [])
    suggestions = analysis.get("resume_suggestions", [])
    summary_points = strengths[:3] or relevant_lines[:3] or ["请基于原简历补充与目标岗位最相关的真实经历。"]

    lines = [
        "# 定制简历草稿",
        "",
        "> 免责声明：本工具仅用于简历表达优化，不应虚构经历。请你逐条确认公司、岗位、时间、学校、奖项、项目和数据真实性。",
        "",
        "## 个人简介",
    ]
    lines.extend([_to_bullet(point) for point in summary_points])
    lines.append("")
    for section in [
        _format_section("教育背景", education),
        _format_section("实习经历", internships or relevant_lines[:5]),
        _format_section("项目经历", projects or relevant_lines[5:10]),
        _format_section("技能", skills),
        _format_section("其他经历", other),
    ]:
        lines.extend(section)
        lines.append("")

    lines.extend([
        "## 针对 JD 的表达优化建议",
    ])
    if suggestions:
        lines.extend([_to_bullet(item) for item in suggestions])
    else:
        lines.append("- 使用“动作 + 方法/工具 + 结果”的结构改写已有经历，但不要新增未发生的结果或指标。")
    lines.extend([
        "",
        "## 使用前人工核对清单",
        "- 是否所有经历都来自原始简历或真实经历？",
        "- 是否删除了无法证明的数据、奖项或成果？",
        "- 是否优先突出了与 JD 相关的能力？",
    ])
    return "\n".join(lines)


def export_resume_files(markdown_text: str, output_dir: Path = OUTPUTS_DIR) -> dict[str, Path]:
    """将定制简历导出为 Markdown 和 DOCX，并返回文件路径。"""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    target_dir = output_dir / timestamp
    target_dir.mkdir(parents=True, exist_ok=True)

    markdown_path = target_dir / "tailored_resume.md"
    markdown_path.write_text(markdown_text, encoding="utf-8")

    docx_path = target_dir / "tailored_resume.docx"
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("当前环境未安装 python-docx，请运行 pip install -r requirements.txt 后再导出 DOCX。") from exc

    document = Document()
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            document.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("> "):
            document.add_paragraph(stripped[2:])
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
        else:
            document.add_paragraph(stripped)
    document.save(docx_path)

    return {"markdown": markdown_path, "docx": docx_path, "directory": target_dir}
